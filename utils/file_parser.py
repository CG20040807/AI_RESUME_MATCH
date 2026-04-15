from docx import Document
from zipfile import BadZipFile


def parse_docx(file):
    try:
        file.seek(0)
        doc = Document(file)

        content = []

        for p in doc.paragraphs:
            if p.text.strip():
                content.append(p.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    content.append(" | ".join(row_text))

        return "\n".join(content)

    except BadZipFile:
        return "【错误】文件损坏或非标准docx"

    except Exception as e:
        return f"【解析失败】{str(e)}"
