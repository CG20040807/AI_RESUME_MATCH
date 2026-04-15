import io
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def export(results, summary, job_title, jd, criteria):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Talent Assessment Report")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_paragraph("")

    doc.add_heading("岗位名称", level=1)
    doc.add_paragraph(job_title)

    doc.add_heading("岗位JD", level=1)
    doc.add_paragraph(jd)

    doc.add_heading("岗位评估标准", level=1)
    doc.add_paragraph(criteria if criteria.strip() else "未填写")

    doc.add_heading("综合总结", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("候选人排名", level=1)
    for idx, r in enumerate(results, start=1):
        doc.add_paragraph(f"{idx}. {r['name']} - {r['score']}分 - {r.get('recommendation', '未提取')}")

    for idx, r in enumerate(results, start=1):
        doc.add_page_break()
        doc.add_heading(f"候选人详情：{idx}. {r['name']}", level=1)
        doc.add_paragraph(r["analysis"])

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio
  
