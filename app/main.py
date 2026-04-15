import sys
from pathlib import Path

# ================== 🚨 自动定位 repo root（最稳版本） ==================

CURRENT = Path(__file__).resolve()

# main.py 在 app/ 下，所以往上找两层才是 repo root
ROOT_DIR = CURRENT.parents[1]

# 兜底：防止部署环境变化
if not (ROOT_DIR / "core").exists():
    ROOT_DIR = CURRENT.parents[2]

if not (ROOT_DIR / "core").exists():
    raise RuntimeError("Cannot locate project root (core/ not found)")

sys.path.insert(0, str(ROOT_DIR))
# ================== 项目模块（必须在路径修复后导入） ==================

from core.analyzer import analyze
from core.scorer import extract_score
from core.ranker import rank_candidates
from core.summarizer import summarize

from utils.file_parser import parse_docx
from utils.text_cleaner import clean_text

# Word导出（容错）
try:
    from utils.docx_exporter import export_to_docx
except Exception:
    export_to_docx = None

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

import streamlit as st

st.write("ROOT_DIR:", ROOT_DIR)
st.write("CORE EXISTS:", (ROOT_DIR / "core").exists())
# ================== 页面配置 ==================

st.set_page_config(
    page_title="AI Talent Assessment System",
    layout="wide"
)

st.markdown("""
<style>
.stApp {
    background: linear-gradient(180deg, #f7fbff 0%, #ffffff 30%, #f8fbf8 100%);
}

.main-title {
    font-size: 2.2rem;
    font-weight: 800;
    color: #1f3b57;
}

.sub-title {
    font-size: 1rem;
    color: #6b7280;
    margin-bottom: 10px;
}
</style>
""", unsafe_allow_html=True)

st.markdown('<div class="main-title">AI Talent Assessment System</div>', unsafe_allow_html=True)
st.markdown('<div class="sub-title">批量简历分析 · 排名 · 总结 · 报告导出</div>', unsafe_allow_html=True)


# ================== Session State ==================

if "results" not in st.session_state:
    st.session_state.results = []

if "summary" not in st.session_state:
    st.session_state.summary = ""

if "word_bytes" not in st.session_state:
    st.session_state.word_bytes = None


# ================== Word fallback ==================

def build_docx(job_title, jd, criteria, results, summary):
    doc = Document()

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI Talent Assessment Report")
    run.bold = True
    run.font.size = Pt(18)

    doc.add_heading("岗位名称", level=1)
    doc.add_paragraph(job_title)

    doc.add_heading("岗位JD", level=1)
    doc.add_paragraph(jd)

    doc.add_heading("评估标准", level=1)
    doc.add_paragraph(criteria)

    doc.add_heading("总结", level=1)
    doc.add_paragraph(summary)

    for i, r in enumerate(results):
        doc.add_page_break()
        doc.add_heading(f"{i+1}. {r['name']}（{r['score']}分）", level=1)
        doc.add_paragraph(r["analysis"])

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer


# ================== 输入区 ==================

col1, col2 = st.columns([1.2, 1])

with col1:
    job_title = st.text_input("岗位名称")
    jd = st.text_area("岗位JD", height=180)
    criteria = st.text_area("评估标准", height=180)

with col2:
    uploaded_files = st.file_uploader(
        "上传简历（.docx）",
        type=["docx"],
        accept_multiple_files=True
    )


# ================== 主流程 ==================

if st.button("开始分析"):

    # ---------- 输入校验 ----------
    if not job_title.strip():
        st.warning("请输入岗位名称")
        st.stop()

    if not jd.strip():
        st.warning("请输入岗位JD")
        st.stop()

    if not criteria.strip():
        st.warning("请输入评估标准")
        st.stop()

    if not uploaded_files:
        st.warning("请上传简历")
        st.stop()

    # ---------- 初始化 ----------
    results = []
    total = len(uploaded_files)
    progress = st.progress(0)

    # ---------- 批量处理 ----------
    for i, file in enumerate(uploaded_files):

        raw_text = parse_docx(file)

        if not raw_text:
            results.append({
                "name": file.name,
                "analysis": "解析失败",
                "score": 0,
                "recommendation": "无法评估"
            })
            continue

        text = clean_text(raw_text)

        # ===== 核心分析 =====
        analysis = analyze(job_title, jd, criteria, text)
        score = extract_score(analysis)

        # ===== 推荐提取 =====
        rec = "未提取"
        m = re.search(r"推荐建议[:：]\s*(.*)", analysis, re.S)
        if m:
            rec = m.group(1).strip()

        results.append({
            "name": file.name,
            "analysis": analysis,
            "score": score,
            "recommendation": rec
        })

        progress.progress((i + 1) / total)

    # ---------- 排序 + 总结 ----------
    ranked = rank_candidates(results)
    summary = summarize(job_title, jd, criteria, ranked)

    # ---------- Word 导出 ----------
    if export_to_docx:
        try:
            word_file = export_to_docx(job_title, jd, criteria, ranked, summary)
        except Exception:
            word_file = build_docx(job_title, jd, criteria, ranked, summary)
    else:
        word_file = build_docx(job_title, jd, criteria, ranked, summary)

    # ---------- 存入 session ----------
    st.session_state.results = ranked
    st.session_state.summary = summary
    st.session_state.word_bytes = word_file


# ================== 展示区 ==================

if st.session_state.results:

    ranked = st.session_state.results
    summary = st.session_state.summary

    top = ranked[0]

    c1, c2, c3 = st.columns(3)

    c1.metric("最佳候选人", top["name"])
    c2.metric("最高分", top["score"])
    c3.metric("推荐意见", top["recommendation"])

    st.markdown("## 排名结果")

    for i, r in enumerate(ranked):
        with st.expander(f"{i+1}. {r['name']}（{r['score']}分）"):
            st.write(r["analysis"])

    st.markdown("## 总结")
    st.write(summary)

    if st.session_state.word_bytes:
        st.download_button(
            "下载Word报告",
            st.session_state.word_bytes,
            file_name="AI_Talent_Report.docx"
        )
